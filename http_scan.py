import os
import re
import xml.etree.ElementTree as Et
import xlsxwriter
from config import  Config
import matplotlib.pyplot as plt
import numpy as np
import tempfile
from docx import Document
from docx.oxml.shared import qn, OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm
from docx.shared import Pt, Inches

class HttpScanner:
    def __init__(self):
        self._hostname = "Hacky"
        self.path_zap = "report_zap"
        self.e_data = {}

    def parse_xml(self):
        data = {}
        for filename in os.listdir(self.path_zap):
            if filename.endswith('.xml'):
                # parse the XML
                tree = Et.parse(os.path.join(self.path_zap, filename))
                root = tree.getroot()

                site = root.find('site')
                site_data = {'name': site.attrib.get('name'), 'host': site.attrib.get('host'),
                             'port': site.attrib.get('port'), 'ssl': site.attrib.get('ssl')}
                data['site'] = site_data

                alerts = site.find('alerts')
                alert_items = []
                for alert_item in alerts.findall('alertitem'):
                    item_data = {'pluginid': alert_item.findtext('pluginid'),
                                 'alertRef': alert_item.findtext('alertRef'), 'alert': alert_item.findtext('alert'),
                                 'name': alert_item.findtext('name'), 'riskcode': alert_item.findtext('riskcode'),
                                 'confidence': alert_item.findtext('confidence'),
                                 'riskdesc': alert_item.findtext('riskdesc'),
                                 'confidencedesc': alert_item.findtext('confidencedesc'),
                                 'desc': alert_item.findtext('desc')}

                    instances = alert_item.find('instances')
                    instance_data = []
                    for instance in instances.findall('instance'):
                        instance_data.append({
                            'uri': instance.findtext('uri'),
                            'method': instance.findtext('method'),
                            'param': instance.findtext('param'),
                            'attack': instance.findtext('attack'),
                            'evidence': instance.findtext('evidence'),
                            'otherinfo': instance.findtext('otherinfo')
                        })
                    item_data['instances'] = instance_data

                    item_data['count'] = alert_item.findtext('count')
                    item_data['solution'] = alert_item.findtext('solution')
                    item_data['otherinfo'] = alert_item.findtext('otherinfo')
                    item_data['reference'] = alert_item.findtext('reference')
                    item_data['cweid'] = alert_item.findtext('cweid')
                    item_data['wascid'] = alert_item.findtext('wascid')
                    item_data['sourceid'] = alert_item.findtext('sourceid')

                    alert_items.append(item_data)

                data['alerts'] = alert_items

                # Extract host, port, and param with the same pluginid
                extracted_data = {}
                for item_data in alert_items:
                    pluginid = item_data['pluginid']
                    instances = item_data['instances']

                    if pluginid not in extracted_data:
                        extracted_data[pluginid] = []

                    for instance in instances:
                        host = data['site']['host']
                        port = data['site']['port']
                        param = instance['uri']


                        extracted_data[pluginid].append({'host': host, 'port': port, 'uri': param})
                self.e_data = extracted_data


            return data

    def remove_tags(self,string):
        string = string.replace("<p>", "")
        string = string.replace("</p>", "")
        return string

    def remove_tags2(self,string):
        string = string.replace("<p>", " ")
        string = string.replace("</p>", " ")
        return string

    def create_excel_sheets(self):
        data = self.parse_xml()

        workbook = xlsxwriter.Workbook("ReportOwasp.xlsx")
        workbook.set_properties({
            'title': "Report Owasp",
            'subject': 'Report Owasp',
            'category': 'report',
            'keywords': 'Owasp, report'})
        # ====================
        # FORMATTING
        # ====================

        # Define formatting styles
        workbook.formats[0].set_font_name('Tahoma')

        format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                          'font_color': '#183868', 'bold': True,
                                                          'align': 'center', 'valign': 'vcenter', 'border': 1})
        format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                                   'font_color': 'white', 'bold': True,
                                                   'align': 'center', 'valign': 'vcenter',
                                                   'border': 1,
                                                   'bg_color': '#183868'})
        format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                                  'align': 'left', 'valign': 'top',
                                                  'border': 1, 'text_wrap': 1})
        format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                                   'align': 'center', 'valign': 'top'})

        """
        Write the extracted information to an Excel file.
        """


        # ====================
        def __row_height(text, width):
            return (max((len(text) // width), text.count('\n')) + 1) * 15

        # ====================
        # TABLE OF CONTENTS
        # ====================
        sheet_name = "TOC"
        ws_toc = workbook.add_worksheet(sheet_name)
        ws_toc.set_tab_color('#183868')

        ws_toc.set_column("A:A", 7)
        ws_toc.set_column("B:B", 5)
        ws_toc.set_column("C:C", 70)
        ws_toc.set_column("D:D", 15)
        ws_toc.set_column("E:E", 50)
        ws_toc.set_column("F:F", 7)

        ws_toc.merge_range("B2:D2", "TABLE OF CONTENTS", format_sheet_title_content)
        ws_toc.write("B3", "No.", format_table_titles)
        ws_toc.write("C3", "Vulnerability", format_table_titles)
        ws_toc.write("D3", "CWE", format_table_titles)


        for index, instance in enumerate(data['alerts'],1):
            vuln_name = instance.get('name')
            desc = instance.get('desc')
            solution = instance.get('solution')
            risk = instance.get('riskdesc')
            severity = risk.split()[0]
            cwe = instance.get('cweid')
            reference = instance.get('reference')


            name = re.sub(r"[\[\]\\\'\"&@#():*?/]", "", vuln_name)
            if len(name) > 27:
                name = "{}..{}".format(name[0:15], name[-10:])
            name = "{:03X}_{}".format(index, name)
            ws_vuln = workbook.add_worksheet(name)
            ws_vuln.set_tab_color(Config.colors()[severity.lower()])

            # --------------------
            # TABLE OF CONTENTS
            # --------------------
            ws_toc.write("B{}".format(index + 3), "{:03X}".format(index), format_table_cells)
            ws_toc.write_url("C{}".format(index + 3), "internal:'{}'!A1".format(name), format_table_cells,
                             string=vuln_name)
            ws_toc.write("D{}".format(index + 3), "{}".format(cwe))
            ws_vuln.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), index + 3),
                              format_align_center,
                              string="<< TOC")
            ws_toc.set_row(index + 3, __row_height(name, 150), None)

            ws_vuln.set_column("A:A", 7, format_align_center)
            ws_vuln.set_column("B:B", 20, format_align_center)
            ws_vuln.set_column("C:C", 20, format_align_center)
            ws_vuln.set_column("D:D", 50, format_align_center)
            ws_vuln.set_column("E:E", 15, format_align_center)
            ws_vuln.set_column("F:F", 15, format_align_center)
            ws_vuln.set_column("G:G", 20, format_align_center)
            ws_vuln.set_column("H:H", 7, format_align_center)
            content_width = 120

            ws_vuln.write('B2', "Title", format_table_titles)
            ws_vuln.merge_range("C2:G2", vuln_name, format_sheet_title_content)
            ws_vuln.set_row(1, __row_height(vuln_name, content_width), None)

            ws_vuln.write('B3', "Description", format_table_titles)
            ws_vuln.merge_range("C3:G3", self.remove_tags(desc), format_table_cells)
            ws_vuln.set_row(2, __row_height(desc, content_width), None)

            ws_vuln.write('B4', "Solution", format_table_titles)
            ws_vuln.merge_range("C4:G4", self.remove_tags(solution), format_table_cells)
            ws_vuln.set_row(3, __row_height(solution, content_width), None)

            ws_vuln.write('B5', "CWE", format_table_titles)
            cvss = int(cwe)
            if cvss >= 0:
                ws_vuln.merge_range("C5:G5", "{:.1f}".format(cvss), format_table_cells)
            else:
                ws_vuln.merge_range("C5:G5", "{}".format("No CWE"), format_table_cells)


            ws_vuln.write('B6', "Level", format_table_titles)
            ws_vuln.merge_range("C6:G6", severity.capitalize(), format_table_cells)
            ws_vuln.write('B7', "References", format_table_titles)
            ws_vuln.merge_range("C7:G7", self.remove_tags2(reference) + "\n", format_table_cells)
            ws_vuln.set_row(10, __row_height(reference, content_width), None)

            ws_vuln.write('C9', "IP", format_table_titles)
            ws_vuln.write('D9', "Host name", format_table_titles)
            ws_vuln.write('E9', "Port number", format_table_titles)


            for j, (id, host) in enumerate(self.e_data.items(),10):
                if id == instance.get('pluginid'):
                    for i in range(len(host)):
                        ip = host[i]['host']
                        port = host[i]['port']
                        uri = host[i]['uri']
                        ws_vuln.write("C{}".format(10 + i), ip)
                        ws_vuln.write("E{}".format(10 + i), port)
                        ws_vuln.write("D{}".format(10 + i), uri)




        workbook.close()


    def create_word(self):
        data = self.parse_xml()

        # ====================
        # DOCUMENT PROPERTIES
        # ====================
        document = Document('src/owasp-template.docx')

        doc_prop = document.core_properties
        doc_prop.title = "Report"
        doc_prop.category = "Report"

        document.add_paragraph('Report Di Sicurezza', style='Title')

        # ====================
        # TABLE OF CONTENTS
        # ====================
        document.add_paragraph('Table of Contents', style='Heading 1')

        par = document.add_paragraph()
        run = par.add_run()
        fld_char = OxmlElement('w:fldChar')  # creates a new element
        fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instr_text.text = r'TOC \h \z \t "OV-H1toc;1;OV-H2toc;2;OV-H3toc;3;OV-Finding;3"'

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        fld_char3 = OxmlElement('w:t')
        fld_char3.text = "# Right-click to update field. #"
        fld_char2.append(fld_char3)

        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)

        document.add_page_break()


        for index, instance in enumerate(data['alerts'], 1):
            vuln_name = instance.get('name')
            desc = instance.get('desc')
            solution = instance.get('solution')
            risk = instance.get('riskdesc')
            severity = risk.split()[0]
            cwe = instance.get('cweid')
            reference = instance.get('reference')
            # --------------------
            # GENERAL
            # --------------------

            title = "{} {}".format(index, vuln_name)
            document.add_paragraph(title, style='Heading 3').paragraph_format.page_break_before = True

            table_vuln = document.add_table(rows=6, cols=2, style="Table Grid")
            table_vuln.autofit = False

            # COLOR
            # --------------------
            col_cells = table_vuln.rows[0].cells
            # col_cells[0].merge(col_cells[1])
            color_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), Config.colors()[severity.lower()][1:]))
            col_cells[1]._tc.get_or_add_tcPr().append(color_fill)

            # Crea un elemento OxmlElement per la sfumatura azzurra
            blue_shading = OxmlElement('w:shd')
            blue_shading.set(qn('w:fill'), '42a5f5')  # Colore azzurro
            blue_shading1 = OxmlElement('w:shd')
            blue_shading1.set(qn('w:fill'), '42a5f5')
            blue_shading2 = OxmlElement('w:shd')
            blue_shading2.set(qn('w:fill'), '42a5f5')
            blue_shading3 = OxmlElement('w:shd')
            blue_shading3.set(qn('w:fill'), '42a5f5')
            blue_shading4 = OxmlElement('w:shd')
            blue_shading4.set(qn('w:fill'), '42a5f5')
            blue_shading5 = OxmlElement('w:shd')
            blue_shading5.set(qn('w:fill'), '42a5f5')
            table_vuln.rows[0].cells[0]._tc.get_or_add_tcPr().append(blue_shading)
            table_vuln.rows[1].cells[0]._tc.get_or_add_tcPr().append(blue_shading1)
            table_vuln.rows[2].cells[0]._tc.get_or_add_tcPr().append(blue_shading2)
            table_vuln.rows[3].cells[0]._tc.get_or_add_tcPr().append(blue_shading3)
            table_vuln.rows[4].cells[0]._tc.get_or_add_tcPr().append(blue_shading4)
            table_vuln.rows[5].cells[0]._tc.get_or_add_tcPr().append(blue_shading5)

            for col_cell in col_cells:
                col_cell.width = Cm(0.42)

            # TABLE HEADERS
            # --------------------
            hdr_cells = table_vuln.columns[0].cells
            hdr_cells[0].paragraphs[0].add_run('Severità').bold = True
            hdr_cells[1].paragraphs[0].add_run('Description').bold = True
            hdr_cells[2].paragraphs[0].add_run('References').bold = True
            hdr_cells[3].paragraphs[0].add_run('Remediation').bold = True
            hdr_cells[4].paragraphs[0].add_run('CWE').bold = True

            for hdr_cell in hdr_cells:
                hdr_cell.width = Cm(3.58)

            # FIELDS
            # --------------------

            txt_cells = table_vuln.columns[1].cells
            txt_cells[0].text = severity
            txt_cells[1].text = self.remove_tags(desc)
            txt_cells[2].text = self.remove_tags2(reference)
            txt_cells[3].text = self.remove_tags(solution)
            txt_cells[4].text = cwe

            for txt_cell in txt_cells:
                txt_cell.width = Cm(12.50)

            # VULN HOSTS
            # --------------------
            para = document.add_paragraph('Vulnerable hosts', style="Heading 4")
            para.paragraph_format.space_before = Inches(0.8)

            # add coloumn for result per port and resize columns
            for j, (id, host) in enumerate(self.e_data.items(), 1):
                if id == instance.get('pluginid'):
                    for i in range(len(host)):
                        ip = host[i]['host']
                        port = host[i]['port']
                        table_hosts = document.add_table(cols=3, rows=(len(ip) + 1), style="Table Grid")

                        col_cells = table_hosts.columns[0].cells
                        for col_cell in col_cells:
                            col_cell.width = Cm(3.2)

                        col_cells = table_hosts.columns[1].cells
                        for col_cell in col_cells:
                            col_cell.width = Cm(3.2)

                        col_cells = table_hosts.columns[2].cells
                        for col_cell in col_cells:
                            col_cell.width = Cm(1.6)

                        hdr_cells = table_hosts.rows[0].cells
                        hdr_cells[0].paragraphs[0].add_run('Host name').bold = True
                        hdr_cells[1].paragraphs[0].add_run('IP').bold = True
                        hdr_cells[2].paragraphs[0].add_run('Port number').bold = True


                        cells = table_hosts.rows[j].cells
                        cells[0].text = " "
                        cells[1].text =  ip if ip else "-"
                        if port and port is not None:
                            cells[2].text = "-" if port == 0 else str(port)
                        else:
                            cells[2].text = "No port info"

                            # Add visible lines to the table
                            table_vuln.style = 'Table Grid'

        output_file = 'owasp_report.docx'

        document.save(output_file)

    def get_path(self):
        return self.path_zap

